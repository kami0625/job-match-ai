"""文档解析边界单元测试

- DOCX 正常解析(段落+表格)
- 非法扩展名拒绝
- 空 docx 拒绝
不依赖外部服务(生成临时文件)。
"""
import sys
from pathlib import Path

import pytest
from docx import Document as DocxDocument

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from app.utils.common_tools import AppError
from app.utils.doc_parser import parse_document, parse_docx


def _make_docx(path: Path, text: str = "个人简历 张三 本科", with_table: bool = False) -> Path:
    doc = DocxDocument()
    doc.add_paragraph(text)
    if with_table:
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "技能"
        table.cell(0, 1).text = "Java"
        table.cell(1, 0).text = "项目"
        table.cell(1, 1).text = "订单系统"
    doc.save(str(path))
    return path


def test_parse_docx_normal(tmp_path):
    p = _make_docx(tmp_path / "resume.docx")
    result = parse_document(str(p))
    assert result["text"]
    assert result["file_name"] == "resume.docx"
    assert "张三" in result["text"]


def test_parse_docx_with_table(tmp_path):
    """表格内容也应被提取"""
    p = _make_docx(tmp_path / "t.docx", with_table=True)
    text = parse_docx(str(p))
    assert "Java" in text
    assert "订单系统" in text


def test_parse_unsupported_ext(tmp_path):
    p = tmp_path / "file.txt"
    p.write_text("hello", encoding="utf-8")
    with pytest.raises(AppError):
        parse_document(str(p))


def test_parse_empty_docx(tmp_path):
    """无内容 docx 拒绝"""
    p = _make_docx(tmp_path / "empty.docx", text="")
    with pytest.raises(AppError):
        parse_document(str(p))


def test_parse_missing_file(tmp_path):
    with pytest.raises(Exception):
        parse_document(str(tmp_path / "not_exists.docx"))
