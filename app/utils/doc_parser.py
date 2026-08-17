"""PDF / Word 文档解析工具

能力：
1. PDF 文本提取（PyMuPDF，含页码信息）
2. Word 文本提取（python-docx，含表格内容）
3. 按扩展名自动路由解析
4. 文本语义分块（段落合并 + 字符切割 + 块间重叠）
"""

import re
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.config import CHUNK_OVERLAP, CHUNK_SIZE
from app.utils.common_tools import AppError, get_logger

logger = get_logger("doc_parser")

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def parse_pdf(file_path: str) -> list:
    """解析 PDF 文件

    Args:
        file_path: PDF 文件路径
    Returns:
        页面文本列表 [{"page_num": 1, "text": "..."}, ...]（仅含非空页）
    """
    pages = []
    try:
        with fitz.open(file_path) as doc:
            for page_index, page in enumerate(doc, start=1):
                text = page.get_text("text").strip()
                if text:
                    pages.append({"page_num": page_index, "text": text})
    except Exception as exc:
        logger.error("PDF 解析失败 path=%s error=%s", file_path, exc)
        raise AppError(f"PDF 解析失败：{exc}") from exc
    if not pages:
        raise AppError("PDF 中未提取到文本内容")
    return pages


def parse_docx(file_path: str) -> str:
    """解析 Word 文档

    Args:
        file_path: docx 文件路径
    Returns:
        拼接后的纯文本（段落 + 表格）
    """
    try:
        doc = DocxDocument(file_path)
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
    except Exception as exc:
        logger.error("Word 解析失败 path=%s error=%s", file_path, exc)
        raise AppError(f"Word 解析失败：{exc}") from exc
    if not text.strip():
        raise AppError("Word 文档中未提取到文本内容")
    return text


def parse_document(file_path: str) -> dict:
    """按扩展名自动解析文档

    Args:
        file_path: 支持 .pdf / .docx
    Returns:
        {"file_name": 文件名, "text": 全文, "pages": [{"page_num","text"}] 或 None}
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise AppError(f"不支持的文件类型：{ext}，仅支持 PDF / DOCX", code=400)
    if ext == ".pdf":
        pages = parse_pdf(str(path))
        text = "\n".join(p["text"] for p in pages)
        return {"file_name": path.name, "text": text, "pages": pages}
    text = parse_docx(str(path))
    return {"file_name": path.name, "text": text, "pages": None}


def split_text(
    text: str,
    chunk_size: Optional[int] = None,
    chunk_overlap: Optional[int] = None,
) -> list:
    """文本语义分块

    策略：按空行切段落，段落合并至目标大小；超长段落按字符切割；
      块间补充重叠片段，保留上下文连贯性。

    Args:
        text: 原始文本
        chunk_size: 单块最大字符数，默认取全局配置
        chunk_overlap: 块间重叠字符数，默认取全局配置
    Returns:
        分块字符串列表
    """
    size = chunk_size or CHUNK_SIZE
    overlap = chunk_overlap or CHUNK_OVERLAP
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list = []
    buffer = ""
    for para in paragraphs:
        while len(para) > size:
            if buffer:
                chunks.append(buffer)
                buffer = ""
            chunks.append(para[:size])
            para = para[size - overlap:]
        if len(buffer) + len(para) <= size:
            buffer = f"{buffer}\n{para}" if buffer else para
        else:
            if buffer:
                chunks.append(buffer)
            buffer = para
    if buffer:
        chunks.append(buffer)
    # 补充块间重叠，增强检索连续性
    merged = []
    for i, chunk in enumerate(chunks):
        if i > 0 and overlap > 0:
            chunk = chunks[i - 1][-overlap:] + chunk
        merged.append(chunk)
    return [c for c in merged if c.strip()]
